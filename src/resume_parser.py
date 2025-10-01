import PyPDF2
import pdfplumber
from docx import Document
import re
import io

class ResumeParser:
    """Class to handle resume parsing from PDF and DOCX files"""
    
    def __init__(self):
        pass
    
    def extract_text_from_pdf(self, pdf_file):
        """Extract text from PDF using pdfplumber for better formatting"""
        text = ""
        try:
            # If it's a file-like object (Streamlit upload), use it directly
            if hasattr(pdf_file, 'read'):
                with pdfplumber.open(io.BytesIO(pdf_file.read())) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            else:
                # If it's a file path
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
        except Exception as e:
            print(f"Error extracting PDF text: {str(e)}")
            # Fallback to PyPDF2
            try:
                if hasattr(pdf_file, 'read'):
                    pdf_file.seek(0)  # Reset file pointer
                    reader = PyPDF2.PdfReader(io.BytesIO(pdf_file.read()))
                else:
                    reader = PyPDF2.PdfReader(pdf_file)
                
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            except Exception as e2:
                print(f"Fallback PDF extraction also failed: {str(e2)}")
                return ""
        
        return text
    
    def extract_text_from_docx(self, docx_file):
        """Extract text from Word document"""
        text = ""
        try:
            if hasattr(docx_file, 'read'):
                # Streamlit uploaded file
                doc = Document(io.BytesIO(docx_file.read()))
            else:
                # File path
                doc = Document(docx_file)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract hyperlinks from the document
            # Many resumes have LinkedIn as a hyperlink
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.element.tag.endswith('hyperlink') or 'hyperlink' in str(run.element.xml):
                        # Try to extract hyperlink URL
                        pass
            
            # Try to extract from document relationships (where hyperlinks are stored)
            try:
                for rel in doc.part.rels.values():
                    if "hyperlink" in rel.reltype:
                        url = rel.target_ref
                        if 'linkedin' in url.lower():
                            text += f"\n{url}\n"
            except:
                pass
                
        except Exception as e:
            print(f"Error extracting DOCX text: {str(e)}")
            return ""
        
        return text
    
    def clean_text(self, text):
        """Clean and normalize text while preserving important patterns like URLs, emails, and phone numbers"""
        if not text:
            return ""
        
        # First, protect URLs and email addresses by temporarily replacing them
        url_pattern = r'https?://[^\s]+'
        urls = re.findall(url_pattern, text)
        url_placeholders = {}
        for i, url in enumerate(urls):
            placeholder = f"__URL_{i}__"
            url_placeholders[placeholder] = url
            text = text.replace(url, placeholder)
        
        # Protect LinkedIn URLs without protocol
        linkedin_pattern = r'(?:www\.)?linkedin\.com/(?:in|pub)/[\w\-]+'
        linkedin_urls = re.findall(linkedin_pattern, text, re.IGNORECASE)
        for i, url in enumerate(linkedin_urls):
            placeholder = f"__LINKEDIN_{i}__"
            url_placeholders[placeholder] = url
            text = text.replace(url, placeholder)
        
        # Protect email addresses
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        email_placeholders = {}
        for i, email in enumerate(emails):
            placeholder = f"__EMAIL_{i}__"
            email_placeholders[placeholder] = email
            text = text.replace(email, placeholder)
        
        # Protect phone numbers
        phone_patterns = [
            r'\+\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
            r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\d{3}\.\d{3}\.\d{4}',
            r'\d{3}-\d{3}-\d{4}',
            r'\b\d{10}\b',
        ]
        phone_placeholders = {}
        for i, pattern in enumerate(phone_patterns):
            phones = re.findall(pattern, text)
            for j, phone in enumerate(phones):
                placeholder = f"__PHONE_{i}_{j}__"
                phone_placeholders[placeholder] = phone
                text = text.replace(phone, placeholder)
        
        # Now clean the text
        # Remove extra whitespace but keep single spaces
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep important ones
        # Keep: letters, numbers, spaces, @, ., -, _, +, (, ), and our placeholders
        text = re.sub(r'[^\w\s@.\-_+()]', ' ', text)
        
        # Remove multiple spaces
        text = re.sub(r' +', ' ', text)
        
        # Restore URLs
        for placeholder, url in url_placeholders.items():
            text = text.replace(placeholder, url)
        
        # Restore emails
        for placeholder, email in email_placeholders.items():
            text = text.replace(placeholder, email)
        
        # Restore phone numbers
        for placeholder, phone in phone_placeholders.items():
            text = text.replace(placeholder, phone)
        
        return text.strip()
    
    def parse_resume(self, uploaded_file):
        """Main parsing function for Streamlit uploaded files"""
        try:
            file_extension = uploaded_file.name.split('.')[-1].lower()
            
            if file_extension == 'pdf':
                text = self.extract_text_from_pdf(uploaded_file)
            elif file_extension in ['docx', 'doc']:
                text = self.extract_text_from_docx(uploaded_file)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Debug: Print raw text before cleaning (you can remove this later)
            print("=== RAW TEXT SAMPLE ===")
            print(text[:500])
            print("======================")
            
            cleaned_text = self.clean_text(text)
            
            # Debug: Print cleaned text sample (you can remove this later)
            print("=== CLEANED TEXT SAMPLE ===")
            print(cleaned_text[:500])
            print("===========================")
            
            return cleaned_text
            
        except Exception as e:
            print(f"Error parsing resume: {str(e)}")
            return ""
